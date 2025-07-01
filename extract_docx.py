from docx import Document
from unidecode import unidecode
import os
import re

### REGEX'S ###
re_month = r"(janeiro|fevereiro|mar[cç]o|abril|maio|junho|julho|agosto|setembro|outubro|novembro|dezembro)"
re_find_table_summary = r"ano[\:\-\s]+.*sistema[\:\-\s]+.*municipio[\:\-\s]+.*data[:=\s]+.*"
re_extract_table_summary = r"\b(?:ano|sistema(?:/solu[cç][aã]o alternativa)?|municipio|uts|data)\b[\:\-\s]+.*?(?=\s+\b(?:ano|sistema|municipio|uts|data|$)\b|$)"
re_extract_date = r"(\d{2}\/\d{2}\/\d{4})"

counter = 0

docx_tables = {}
docx_tables_data = []
temp_arr = []

# Load the .docx file
filename = "teste5.docx"
doc = Document(f"./amostragem/{filename}")

# Print each paragraph
for para in doc.paragraphs:
    if re.findall(re_find_table_summary, unidecode(para.text), re.IGNORECASE):
        data = re.findall(re_extract_table_summary, unidecode(para.text), re.IGNORECASE)   
        for i in range(0, len(data), 4):
            counter += 1
            summary_structure = {"ano": "", "sistema": "", "municipio": "", "data": ""}
            summary_structure["ano"] = re.split(r"^\w+[\-\:\s]+", re.split(r"[\:\-](?=\b\w+\b)", data[i], re.IGNORECASE)[0], re.IGNORECASE)[1]
            summary_structure["sistema"] = re.split(r"^\w+[\-\:\s]+", re.split(r"[\:\-](?=\b\w+\b)", data[i+1], re.IGNORECASE)[0], re.IGNORECASE)[1]
            summary_structure["municipio"] = re.split(r"^\w+[\-\:\s]+", re.split(r"[\:\-](?=\b\w+\b)", data[i+2], re.IGNORECASE)[0], re.IGNORECASE)[1]
            summary_structure["data"] = re.findall(re_extract_date, data[i+3], re.IGNORECASE)
            docx_tables[f"cabecalho {counter}"] = { "tabelas": [], "dados_de_cabecalho": summary_structure}
    else:
        print(unidecode(para.text))
print(docx_tables)
counter = 0

for t in doc.tables:    #ativa o algoritmo de leitura de tabelas: lê os dados, separa em tabelas e os guarda em uma variável global
    table_rows = []
    for r in t.rows:
        row_data = [cell.text.strip() for cell in r.cells]
        table_rows.append(row_data)
    temp_arr.append(table_rows)

for t in range(len(temp_arr)):
    temp_months = []
    for r in range(len(temp_arr[t])):
        if re.search(re_month, str(temp_arr[t][r]), re.IGNORECASE):
            temp_months = {cell for cell in temp_arr[t][r] if re.search(re_month, str(cell), re.IGNORECASE)}
            temp_months = list(temp_months)
        if re.search(r"(3.frequ[eê]ncia)", str(temp_arr[t][r]), re.IGNORECASE):
            temp_arr[t] = temp_arr[t][r+1:]
            break
    temp_arr[t] = [temp_months] + [temp_arr[t]]

for t in temp_arr:
    temp_months = t[0]
    temp_arr = [r[0] for r in t[1] if r[0]]
    temp_arr = {key: [] for key in temp_arr}
    counter += 1
    data_structure = {"nome": f"tabela {counter}", "dados": temp_arr}
    for r in range(len(t[1])):
        if t[1][r][0] in temp_arr:
            data_structure["dados"][t[1][r][0]] += [cell for cell in t[1][r][1:]]
        else:
            data_structure["dados"][t[1][r-1][0]] = [data_structure["dados"][t[1][r-1][0]],t[1][r][1:]]
    data_structure["meses"] = temp_months
    docx_tables_data.append(data_structure)
    
decrementer = 0

for x in range(len(docx_tables_data)):
    if "Dezembro" not in docx_tables_data[x]["meses"]:
        docx_tables[f"cabecalho {x - decrementer + 1}"]["tabelas"].append(docx_tables_data[x])
        decrementer += 1
    else:
        docx_tables[f"cabecalho {x - decrementer + 1}"]["tabelas"].append(docx_tables_data[x])

# print(docx_tables["cabecalho 1"])

filename = f"{filename.split(".")[0]}_extraction.txt"

# Write to the file in the current directory
with open(filename, "w", encoding="utf-8") as f:
    f.write(str(docx_tables))

print(f"✅ File '{filename}' has been written to the current directory.")


        
